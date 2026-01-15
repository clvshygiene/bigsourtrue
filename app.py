import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document 
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION 
from docx.oxml.ns import qn, nsdecls 
from docx.oxml import parse_xml 

# --- 頁面設定 ---
st.set_page_config(page_title="校園掃區檢核系統", page_icon="🧹", layout="centered")
st.title("🧹 114-2 校園大掃除檢核系統")

# --- 1. 讀取資料函式 ---
@st.cache_data(ttl=600)
def load_data():
    try:
        # 👇 請確認這裡填寫的是正確的 Google 試算表連結
        google_sheet_url = "https://docs.google.com/spreadsheets/d/1jqpj-DOe1X2cf6cToWmtW19_0FdN3REioa34aXn4boA/edit?usp=sharing"
        
        if "/edit" in google_sheet_url:
            excel_url = google_sheet_url.replace("/edit", "/export?format=xlsx")
            excel_url = excel_url.split("?")[0] + "?format=xlsx"
        else:
            excel_url = google_sheet_url

        all_sheets = pd.read_excel(excel_url, sheet_name=None, dtype=str)
        
        required_sheets = ['班級清單', '地點資料庫', '掃區分配總表', '檢查標準']
        for sheet in required_sheets:
            if sheet not in all_sheets:
                st.error(f"❌ 找不到工作表：「{sheet}」")
                return None, None, None

        df_classes = all_sheets['班級清單']
        df_locations = all_sheets['地點資料庫']
        df_assign = all_sheets['掃區分配總表']
        df_standards = all_sheets['檢查標準']
        
        df_full = pd.merge(df_assign, df_locations, on='地點ID', how='left')
        df_full = pd.merge(df_full, df_classes, left_on='負責班級', right_on='班級代碼', how='left')
        df_full = df_full.dropna(subset=['負責班級'])
        
        return df_classes, df_full, df_standards
        
    except Exception as e:
        st.error(f"❌ 資料讀取失敗！錯誤訊息：{e}")
        return None, None, None

# --- 黑魔法函式：設定儲存格背景顏色 ---
def set_cell_bg(cell, hex_color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# --- 輔助函式：建立簽名區 ---
def add_signature_block(doc):
    doc.add_paragraph("\n") 
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'
    
    for row in sig_table.rows:
        row.height = Cm(2.2) 
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def set_cell_text(cell, text):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    set_cell_text(sig_table.cell(0, 0), " 衛生股長") 
    set_cell_text(sig_table.cell(0, 1), " 衛生糾察")
    set_cell_text(sig_table.cell(1, 0), " 導師簽名")
    set_cell_text(sig_table.cell(1, 1), " 衛生組核章")

    # 底部提醒文字
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12) 
    reminder_text = "各位同學好：打掃完之後，先由班上正副衛生股長檢查以後，再請導師簽名。最後請班上同學打電話至衛生組(分機312)，衛生組將會派衛生糾察到場，最後由衛生糾察檢查確認打勾，「由衛生糾察帶回學務處衛生組」。"
    run = p.add_run(reminder_text)
    run.font.size = Pt(12) 
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 輔助函式：建立任務清單區 ---
def add_task_section(doc, tasks_df, standards_grouped, title_text):
    heading = doc.add_heading(title_text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(20) 
        run.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.color.rgb = RGBColor(0, 0, 0)

    heading.paragraph_format.space_after = Pt(12)

    for index, row in tasks_df.iterrows():
        bldg = str(row['大樓']) if pd.notna(row['大樓']) else ""
        floor = str(row['樓層']) if pd.notna(row['樓層']) else ""
        detail = str(row['詳細位置']) if pd.notna(row['詳細位置']) else ""
        full_name = f"{bldg} {floor} {detail}".strip()
        
        h2 = doc.add_heading(f"📍 {full_name}", level=2)
        h2.paragraph_format.space_before = Pt(18) 
        h2.paragraph_format.space_after = Pt(6)
        
        for run in h2.runs:
            run.font.size = Pt(14)
            run.bold = True
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        note = row['特別注意事項']
        if pd.notna(note) and str(note).strip() != "":
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ 注意：{note}")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        
        check_type = row['檢查類型']
        if check_type in standards_grouped.groups:
            type_df = standards_grouped.get_group(check_type)
            
            # 【修正】改回 3 欄 (子分類, 項目, 確認)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.allow_autofit = False 
            
            # --- 表頭設定 ---
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '子分類'
            hdr_cells[1].text = '檢查項目'
            hdr_cells[2].text = '確認'
            
            # 設定表頭底色
            for cell in hdr_cells:
                set_cell_bg(cell, "D9D9D9")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            
            # 設定欄寬 (總寬約 18.5 cm)
            # 子分類 3.5cm | 項目 13.5cm | 確認 1.5cm
            table.columns[0].width = Cm(3.5)
            table.columns[1].width = Cm(13.5)
            table.columns[2].width = Cm(1.5)
            
            hdr_cells[0].width = Cm(3.5)
            hdr_cells[1].width = Cm(13.5)
            hdr_cells[2].width = Cm(1.5)

            if '子分類' in type_df.columns:
                type_df_sorted = type_df.sort_values(by=['子分類'], na_position='first')
            else:
                type_df_sorted = type_df

            for item_row in type_df_sorted.itertuples():
                row_cells = table.add_row().cells
                row_cells[0].height = Cm(1.0)
                
                # 1. 子分類欄
                sub_cat_text = str(item_row.子分類) if pd.notna(item_row.子分類) else ""
                row_cells[0].text = sub_cat_text
                row_cells[0].width = Cm(3.5)
                row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 子分類置中對齊
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 2. 檢查項目欄
                row_cells[1].text = item_row.檢查細項
                row_cells[1].width = Cm(13.5)
                row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 3. 確認欄
                row_cells[2].width = Cm(1.5)
                row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = row_cells[2].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("□")
                run.font.size = Pt(16)

                # 統一設定內容字型
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text != "□": # 方框維持自己的大小
                                run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        else:
            doc.add_paragraph(f"(未找到類型 {check_type} 的檢查標準)")
            
        doc.add_paragraph("") 

    add_signature_block(doc)

# --- 核心邏輯：生成單一班級的內容 ---
def append_class_content(doc, display_name, tasks_df, standards_grouped):
    df_indoor = tasks_df[tasks_df['檢查類型'] == '內掃教室']
    df_outdoor = tasks_df[tasks_df['檢查類型'] != '內掃教室']

    # 1. 內掃頁
    if not df_indoor.empty:
        add_task_section(doc, df_indoor, standards_grouped, f"{display_name} - 內掃教室")
        
        if not df_outdoor.empty:
            section = doc.add_section(WD_SECTION.ODD_PAGE)
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
    
    # 2. 外掃頁
    if not df_outdoor.empty:
        add_task_section(doc, df_outdoor, standards_grouped, f"{display_name} - 外掃區域")

# --- 主程式 ---
df_classes, df_tasks, df_standards = load_data()

if df_tasks is not None:
    
    st.sidebar.header("📍 班級登入")
    
    if '年級' in df_classes.columns:
        all_grades = sorted(df_classes['年級'].astype(str).unique())
        selected_grade = st.sidebar.selectbox("請選擇年級", all_grades)
        classes_filter = df_classes[df_classes['年級'] == selected_grade]
    else:
        st.error("班級清單缺少「年級」欄位")
        st.stop()
    
    class_options = {
        f"{row['班級代碼']} - {row['顯示名稱']}": row['班級代碼'] 
        for index, row in classes_filter.iterrows()
    }
    
    st.sidebar.markdown("---")
    st.sidebar.header("🖨️ 行政專用：批次列印")
    
    if st.sidebar.button("📥 下載「全校」合併 Word 檔"):
        with st.spinner("正在生成全校表單，請稍候..."):
            doc = Document()
            
            section = doc.sections[0]
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
            
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

            standards_grouped = df_standards.groupby('檢查類型')
            all_classes_sorted = df_classes.sort_values(by=['班級代碼'])
            
            first_class = True
            for idx, class_row in all_classes_sorted.iterrows():
                class_id = class_row['班級代碼']
                class_display = class_row['顯示名稱']
                class_tasks = df_tasks[df_tasks['負責班級'] == class_id]
                
                if not class_tasks.empty:
                    if not first_class:
                        section = doc.add_section(WD_SECTION.ODD_PAGE)
                        section.top_margin = Cm(1.27)
                        section.bottom_margin = Cm(1.27)
                        section.left_margin = Cm(1.27)
                        section.right_margin = Cm(1.27)
                    
                    append_class_content(doc, class_display, class_tasks, standards_grouped)
                    first_class = False

            bio = io.BytesIO()
            doc.save(bio)
            
            st.sidebar.download_button(
                label="✅ 點此下載全校檔案",
                data=bio.getvalue(),
                file_name=f"全校大掃除檢核表_合併檔_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    selected_option = st.sidebar.selectbox("請選擇班級 (個別檢視)", list(class_options.keys()))
    current_class_id = class_options[selected_option]
    
    if " - " in selected_option:
        current_display_name = selected_option.split(" - ")[-1]
    else:
        current_display_name = selected_option

    st.info(f"👋 歡迎 **{current_display_name}**")
    
    my_tasks = df_tasks[df_tasks['負責班級'] == current_class_id]
    standards_grouped = df_standards.groupby('檢查類型')
    
    if not my_tasks.empty:
        st.markdown("### 🖨️ 紙本檢核表下載 (單班)")
        
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        
        append_class_content(doc, current_display_name, my_tasks, standards_grouped)
        
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label=f"📥 下載 {current_display_name} Word 檔",
            data=bio.getvalue(),
            file_name=f"{current_display_name}_大掃除檢核表.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.markdown("---")

    st.markdown("### 📱 數位預覽")
    if my_tasks.empty:
        st.warning("目前無分配掃區。")
    else:
        with st.form(key='preview_form'):
            for index, row in my_tasks.iterrows():
                bldg = str(row['大樓']) if pd.notna(row['大樓']) else ""
                floor = str(row['樓層']) if pd.notna(row['樓層']) else ""
                detail = str(row['詳細位置']) if pd.notna(row['詳細位置']) else ""
                full_name = f"{bldg} {floor} {detail}".strip()
                
                check_type = row['檢查類型']
                note = row['特別注意事項']
                location_id = row['地點ID']
                
                st.subheader(f"📍 {full_name}")
                if pd.notna(note) and str(note).strip() != "":
                    st.warning(f"注意：{note}")
                
                if check_type in standards_grouped.groups:
                    type_df = standards_grouped.get_group(check_type)
                    
                    if '子分類' in type_df.columns:
                        sub_groups = type_df.groupby('子分類', sort=False)
                        for sub_cat, items_df in sub_groups:
                            if pd.notna(sub_cat):
                                st.markdown(f"**🔹 {sub_cat}**")
                            
                            cols = st.columns(2)
                            for idx, item_row in enumerate(items_df.itertuples()):
                                unique_key = f"{current_class_id}_{location_id}_{sub_cat}_{item_row.檢查細項}_{idx}"
                                with cols[idx % 2]:
                                    st.checkbox(item_row.檢查細項, key=unique_key)
                            st.write("")
                    else:
                        for idx, item_row in enumerate(type_df.itertuples()):
                             unique_key = f"{current_class_id}_{location_id}_{item_row.檢查細項}_{idx}"
                             st.checkbox(item_row.檢查細項, key=unique_key)
                
                st.markdown("---")
            
            st.form_submit_button("數位送出 (測試用)")
